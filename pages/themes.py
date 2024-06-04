import altair as alt
import numpy as np
import pandas as pd
import streamlit as st
import os
import altair as alt
import matplotlib.pyplot as plt  
import squarify
import time
import plotly.express as px
from docx import Document  
from docx.shared import Inches  
import io
# Spinner Animation
import json
import requests
import streamlit as st
from streamlit_lottie import st_lottie
# Ewan imports
import vertexai
import json
from vertexai.generative_models import GenerativeModel, ChatSession, SafetySetting, HarmCategory, HarmBlockThreshold, GenerationConfig
import pandas as pd
import concurrent.futures

st.set_page_config(
    page_title="Spirit AI",
    page_icon="images/VF_Icon_RGB_RED.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Used for styling
st.markdown(
    """
    <style>
        [data-testid=stSidebar] [data-testid=stImage]{
            text-align: center;
            display: block;
            margin-left: auto;
            margin-right: auto;
            width: 100%;
        }

        [data-testid=stSidebarUserContent] {
            padding-bottom: 0px;
        }

        .sidebar .sidebar-content {
            display: flex;
            flex-direction: column;
            align-items: center;
        }

        .wrap {
            height: 260px;
            margin: 10px;
            display: flex;
            justify-content: center;
            align-items: center;
        }

        .wrap span {
            align-self: flex-end;
        }

    </style>
    """, unsafe_allow_html=True
)


# Ewan Code


# Safety config
safety_config = [
    SafetySetting(
        category=HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT,
        threshold=HarmBlockThreshold.BLOCK_NONE,
    ),
    SafetySetting(
        category=HarmCategory.HARM_CATEGORY_HATE_SPEECH,
        threshold=HarmBlockThreshold.BLOCK_NONE,
    ),
    SafetySetting(
        category=HarmCategory.HARM_CATEGORY_HARASSMENT,
        threshold=HarmBlockThreshold.BLOCK_NONE,
    ),
]

def generate_json_response(user_comment, model_name="gemini-1.5-flash-preview-0514", temperature=0.2, system_prompt=system_prompt_v4):
    project_id = "vf-uk-aib-dev-ads-cht-lab"
    vertexai.init(project=project_id)

    model = GenerativeModel(
        model_name=model_name,
        system_instruction=[system_prompt],
        safety_settings=safety_config,
        generation_config=GenerationConfig(
            temperature=temperature,
            candidate_count=1
        ),    
    )

    prompt = f"""{user_comment}
    """
    contents = [prompt]
    response = model.generate_content(contents)
    return response.text


def extract_json_multiple(text):
    # Remove code block formatting
    text = text.strip('`')
    if text.startswith('json'):
        text = text[4:].strip()

    json_objects = []
    start_index = 0

    while True:
        try:
            # Look for the next occurrence of '{'
            start_index = text.find('{', start_index)
            if start_index == -1:
                break

            # Find the corresponding closing '}'
            end_index = start_index + 1
            brace_count = 1
            while brace_count > 0 and end_index < len(text):
                if text[end_index] == '{':
                    brace_count += 1
                elif text[end_index] == '}':
                    brace_count -= 1
                end_index += 1

            # Extract the JSON substring
            json_substring = text[start_index:end_index]

            # Try parsing the JSON substring
            json_data = json.loads(json_substring)
            json_objects.append(json_data)

            start_index = end_index
        except json.JSONDecodeError:
            start_index += 1

    return json_objects


def process_comment(row):
    global json_objects_list
    
    user_comment = row["comment"]
    
    #debugging
    print(user_comment)
    
    # Call generate_json_response function
    json_response = generate_json_response(user_comment)
    
    # Call extract_json_multiple function
    json_objects = extract_json_multiple(json_response)
    
    # Append the extracted JSON objects to the list
    json_objects_list.extend(json_objects)
    
    return json_objects


def process_comments_dataframe(df, max_workers=16):
    global json_objects_list
    
    # Create a ThreadPoolExecutor with the desired number of workers
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit the tasks to the executor
        futures = [executor.submit(process_comment, row) for _, row in df.iterrows()]
        
        # Wait for all the tasks to complete
        concurrent.futures.wait(futures)
    
    return json_objects_list


def json_to_dataframe(json_objects_list):
    # Create an empty list to store the flattened JSON objects
    flattened_data = []
    
    # Iterate over each JSON object in the list
    for json_object in json_objects_list:
        # Flatten the JSON object
        flattened_object = {}
        for key, value in json_object.items():
            if isinstance(value, dict):
                # If the value is a nested dictionary, flatten it recursively
                for nested_key, nested_value in value.items():
                    flattened_object[f"{key}_{nested_key}"] = nested_value
            elif isinstance(value, list):
                # If the value is a list, convert it to a string representation
                flattened_object[key] = str(value)
            else:
                flattened_object[key] = value
        
        # Append the flattened JSON object to the list
        flattened_data.append(flattened_object)
    
    # Create a DataFrame from the flattened data
    df = pd.DataFrame(flattened_data)
    
    return df



# Set spinner visbility to False
showSpinner = True

# Sidebar
with st.sidebar:
    # VF Icon
    st.image('images/VF_Icon_RGB_RED.png', width=100)

    # Center-aligned header
    st.markdown("<h3 style='text-align: center;'>Spirit Survey AI Analyser</h3>", unsafe_allow_html=True)

    # Divider
    st.divider()

    # Nav page links
    st.page_link("streamlit_app.py", label="Home")
    st.page_link("pages/themes.py", label="Themes")
    st.page_link("pages/sentiment.py", label="Sentiment", disabled=True)
    st.page_link("http://www.google.com", label="Chat", disabled=True)

    # Center-aligned paragraph
    st.markdown("<div style='text-align: center;' class='wrap'><span style='text-align: center;'>Developed by VBTS UK</span></div>", unsafe_allow_html=True)


# Main page

# Center-aligned header
st.markdown("<h1 style='text-align: center;'>Identify Comment Themes</h1>", unsafe_allow_html=True)

# Center-aligned paragraph
st.markdown("<p style='text-align: center;'>Upload your Spirit Beat comment extract (excel file) to perform a thematic analysis.</p>", unsafe_allow_html=True)

# Spacer
st.markdown("<div style='padding: 10px 5px;'></div>", unsafe_allow_html=True)



# # Treemap
# # Function to save the treemap figure to a BytesIO object  
# def save_treemap_to_bytes(fig):  
#     buf = io.BytesIO()  
#     fig.savefig(buf, format='png')  
#     buf.seek(0)  
#     return buf  
  
# # Function to create a .docx file with the treemap image  
# def create_treemap_docx(fig_buf, filename='treemap.docx'):  
#     # Create a new Document  
#     doc = Document()  
#     doc.add_heading('Treemap', 0)  
  
#     # Add the treemap image to the document  
#     doc.add_picture(fig_buf, width=Inches(6))  
      
#     # Save the document to a BytesIO object  
#     doc_buf = io.BytesIO()  
#     doc.save(doc_buf)  
#     doc_buf.seek(0)  
      
#     return doc_buf


# File Uploader
uploaded_files = st.file_uploader("Choose an .xlsx file", accept_multiple_files=True)

if uploaded_files:
    # Set spinner visbility to True
    showSpinner = True

    for uploaded_file in uploaded_files:
        # Read the file into a Pandas DataFrame  
        input_data = pd.read_excel(uploaded_file, skiprows=2)  # Skip the first two rows

        # Check if 'What is one thing that needs to change in order for you and your team to be more successful?' column exists in the DataFrame
        if 'What is one thing that needs to change in order for you and your team to be more successful?' in input_data.columns:
  
            # Extract the desired column  
            column_data = input_data['What is one thing that needs to change in order for you and your team to be more successful?'] 
    
            # Create a new DataFrame with only one column  
            input_df = pd.DataFrame({'comment': column_data})  

            # Print the new DataFrame  
            # st.write(input_df)

            # PIP ADD AI CODE HERE
            # output_data = pd.DataFrame({  
            #     'ID': [1, 2, 3, 4],  
            #     'Comment': [  
            #         'Our team has been reduced and we are all trying to cover each other as best as possible. We need backfill.',  
            #         'We should connect more to have informal chats, to get to know each other.',  
            #         'A purpose built Lab to test our Design Solutions',  
            #         'More cross training between the team'  
            #     ],  
            #     'subtheme': ['Team size reduction', 'Team cohesion', 'Available resources', 'Training availability'],  
            #     'theme_number': ['', '', '', ''],  
            #     'theme_name': ['Teams', 'Teams', 'Teams', 'Training']  
            # })


            #define a global list to populate with the json
            json_objects_list = []

            #use the ai to automagically populate the json list
            process_comments_dataframe(input_df, max_workers=8)

            #convert to a df
            output_data=json_to_dataframe(json_objects_list)


            # Turn off spinner once data returned
            # PIP REMOVE
            # time.sleep(4)
            if not output_data.empty:
                showSpinner = False

                # Check if 'theme_name' column exists in the DataFrame
                if 'theme' in output_data.columns:
                    # Aggregate the data to count occurrences of each theme
                    theme_counts = output_data['theme'].value_counts().reset_index()
                    theme_counts.columns = ['theme', 'count']

                    # Colors for the treemap
                    colors = ['#9C2AA0', '#5E2750', '#00B0CA', '#007C92', '#A8B400', '#FECB00', '#EB9700']
                    
                    # Create the treemap
                    fig = px.treemap(
                        theme_counts,
                        path=['theme'],
                        values='count',
                        color='theme',
                        color_discrete_sequence=colors
                    )
                    
                    # Plot the chart
                    st.plotly_chart(fig, use_container_width=True)

                    # # Save the figure to a BytesIO object  
                    # fig_buf = save_treemap_to_bytes(fig)  
                        
                    # # Create a .docx file with the treemap image  
                    # doc_buf = create_treemap_docx(fig_buf)  


                    # Save the figure as an image  
                    fig.write_image("treemap.png")  
                    
                    # Create a new Word document  
                    doc = Document()  
                    
                    # Add a heading  
                    doc.add_heading("Treemap Chart", level=1)  
                    
                    # Add the Plotly figure as an image to the document  
                    doc.add_picture("treemap.png", width=Inches(6))  
                    
                    # Save the Word document  
                    doc.save("output.docx")  
                        
                    # Use Streamlit's download button to offer the .docx file for download  
                    # st.download_button(  
                    #     label='Download Treemap as DOCX',  
                    #     data=doc,  
                    #     file_name='treemap.docx',  
                    #     mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'  
                    # )  


                else:
                    st.write("The output file does not contain column 'theme_name'.")

        else:
            st.write("The uploaded file does not contain column 'What is one thing that needs to change in order for you and your team to be more successful?'")



# Load animation JSON file - function
def load_lottiefile(filepath: str):
    with open(filepath, "r") as f:
        return json.load(f)


# Load animation JSON file - call function using URL
lottie_coding = load_lottiefile("images/VF_Living_Speechmark_Rotation_drawn_on_Lottie_Red_300.json")  # replace link to local lottie file

# Add columns to centralise loading spinner
col1, col2, col3 = st.columns(3)

# Add loading spinner
if showSpinner:
    with col2:
        st_lottie(
            lottie_coding,
            width=300
        )


# # Colors for the treemap
# colors = ['#9C2AA0', '#5E2750', '#00B0CA', '#007C92', '#A8B400', '#FECB00', '#EB9700']

# # File Uploader
# uploaded_files_output = st.file_uploader("Temp", accept_multiple_files=True)

# # Check if any files have been uploaded
# if uploaded_files_output:
#     time.sleep(4) # Sleep for 4 seconds
#     showSpinner = False
#     for uploaded_file_output in uploaded_files_output:
#         # Read the file into a Pandas DataFrame
#         output_data = pd.read_excel(uploaded_file_output)
        
#         # Display the filename
#         st.write("Filename:", uploaded_file_output.name)
        
#         # Display the DataFrame
#         st.dataframe(output_data)

#         # Check if 'theme_name' column exists in the DataFrame
#         if 'theme_name' in output_data.columns:
#             # Aggregate the data to count occurrences of each theme
#             theme_counts = output_data['theme_name'].value_counts().reset_index()
#             theme_counts.columns = ['theme_name', 'count']
            
#             # Create the treemap
#             fig = px.treemap(
#                 theme_counts,
#                 path=['theme_name'],
#                 values='count',
#                 color='theme_name',
#                 color_discrete_sequence=colors
#             )
            
#             # Plot the chart
#             st.plotly_chart(fig, use_container_width=True)
#         else:
#             st.write("The uploaded file does not contain a 'theme_name' column.")






# volume = [350, 220, 170, 150, 50]
# labels = ['Do the right thing', 'Pay us more',
#          'Management = BAD', 'We should\n work together',
#          'I dont know']
# color_list = ['#e60000', '#4a4d4e', '#9c2aa0',
#              '#00b0ca', '#eb9800', '#5e2750']
 
# fig, ax = plt.subplots()
# plt.rc('font', size=14)
# squarify.plot(sizes=volume, label=labels,
#              color=color_list)
# plt.axis('off')
# # Display the treemap
# st.pyplot(fig)
 
 






# import numpy as np
# import plotly.express as px

# colors = ['#9C2AA0',  '#5E2750', '#00B0CA', '#007C92', '#A8B400', '#FECB00', '#EB9700']

# # Create distplot with custom bin_size
# fig = px.treemap(
#     names = ["Eve","Cain", "Seth", "Enos", "Noam", "Abel", "Awan", "Enoch", "Azura"],
#     parents = ["", "Eve", "Eve", "Seth", "Seth", "Eve", "Eve", "Awan", "Eve"],
# )

# # Plot!
# st.plotly_chart(fig, use_container_width=True)

